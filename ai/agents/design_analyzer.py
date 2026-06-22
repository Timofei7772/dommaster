"""
DesignAnalyzerAgent — чтение дизайн-проектов (PDF, DWG, DXF, DOCX, Excel).
"""
import logging
import os
from typing import Optional

from ai.agents.base_agent import BaseAgent

logger = logging.getLogger(__name__)

SYSTEM_PROMPT = """
Ты — эксперт по чтению дизайн-проектов и строительной документации.
Проанализируй текст/данные из дизайн-проекта.
Извлеки:
1. Список помещений с размерами (площадь, периметр, высота потолков)
2. Тип отделки каждого помещения (пол, стены, потолок)
3. Указанные материалы
4. Особые требования

Ответ строго JSON:
{
    "project_name": "...",
    "total_area": 0.0,
    "rooms": [
        {
            "name": "Гостиная",
            "area": 25.0,
            "perimeter": 20.0,
            "ceiling_height": 2.7,
            "wall_area": 54.0,
            "floor_finish": "паркет дуб",
            "wall_finish": "декоративная штукатурка",
            "ceiling_finish": "натяжной матовый",
            "has_wet_zone": false,
            "features": ["камин", "ниша под TV"],
            "specified_materials": [
                {"name": "Паркет дуб Coswick", "area": 25.0}
            ]
        }
    ],
    "general_notes": "",
    "confidence": 0.9
}
"""


class DesignAnalyzerAgent(BaseAgent):

    async def analyze(
        self,
        file_path: str,
        file_type: str = "pdf",
    ) -> dict:
        text = await self._extract_text(file_path, file_type)

        if not text:
            logger.warning("Не удалось извлечь текст из %s", file_path)
            return {"error": "Не удалось извлечь данные", "rooms": []}

        result = await self._call_llm(
            SYSTEM_PROMPT,
            f"Содержимое дизайн-проекта:\n\n{text[:15000]}",
        )

        logger.info(
            "DesignAnalyzer: проект=%s, комнат=%d",
            result.get("project_name", "?"),
            len(result.get("rooms", [])),
        )
        return result

    async def _extract_text(self, file_path: str, file_type: str) -> str:
        if not os.path.exists(file_path):
            logger.error("Файл не найден: %s", file_path)
            return ""

        extractors = {
            "pdf": self._extract_from_pdf,
            "docx": self._extract_from_docx,
            "xlsx": self._extract_from_excel,
            "xls": self._extract_from_excel,
            "dwg": self._extract_from_dwg,
            "dxf": self._extract_from_dxf,
        }

        extractor = extractors.get(file_type.lower())
        if not extractor:
            logger.warning("Неподдерживаемый формат: %s", file_type)
            return ""

        try:
            return await extractor(file_path)
        except Exception as e:
            logger.exception("Ошибка извлечения из %s", file_path)
            return ""

    @staticmethod
    async def _extract_from_pdf(file_path: str) -> str:
        try:
            import pdfplumber
            text_parts = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                    # Извлекаем таблицы
                    tables = page.extract_tables()
                    for table in tables:
                        for row in table:
                            text_parts.append(" | ".join(
                                str(cell) for cell in row if cell
                            ))
            return "\n".join(text_parts)
        except ImportError:
            # Fallback на PyMuPDF
            import fitz  # PyMuPDF
            doc = fitz.open(file_path)
            text_parts = [page.get_text() for page in doc]
            doc.close()
            return "\n".join(text_parts)

    @staticmethod
    async def _extract_from_docx(file_path: str) -> str:
        from docx import Document as DocxDocument
        doc = DocxDocument(file_path)
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        # Таблицы
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)

    @staticmethod
    async def _extract_from_excel(file_path: str) -> str:
        import openpyxl
        wb = openpyxl.load_workbook(file_path, data_only=True)
        parts = []
        for sheet in wb.sheetnames:
            ws = wb[sheet]
            parts.append(f"=== Лист: {sheet} ===")
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) for c in row if c is not None]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)

    @staticmethod
    async def _extract_from_dxf(file_path: str) -> str:
        import ezdxf
        doc = ezdxf.readfile(file_path)
        parts = []
        msp = doc.modelspace()
        for entity in msp:
            if entity.dxftype() == "TEXT":
                parts.append(entity.dxf.text)
            elif entity.dxftype() == "MTEXT":
                parts.append(entity.text)
            elif entity.dxftype() == "LINE":
                start = entity.dxf.start
                end = entity.dxf.end
                length = (
                    (end.x - start.x) ** 2
                    + (end.y - start.y) ** 2
                ) ** 0.5
                parts.append(f"Линия: {length:.2f} мм")
        return "\n".join(parts)

    @staticmethod
    async def _extract_from_dwg(file_path: str) -> str:
        """DWG → конвертация в DXF через ODA или чтение через ezdxf."""
        logger.warning("Прямое чтение DWG ограничено, рекомендуется конвертация в DXF")
        # Попробуем как DXF (некоторые версии совместимы)
        try:
            return await DesignAnalyzerAgent._extract_from_dxf(file_path)
        except Exception:
            return ""
