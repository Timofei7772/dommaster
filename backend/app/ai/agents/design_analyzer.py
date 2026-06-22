"""
Агент анализа дизайн-проекта — чтение PDF, DWG, DXF, DOCX, Excel
"""

from typing import Dict, Any
import os
from app.ai.base_agent import BaseAgent
from app.ai.prompts import DESIGN_ANALYZER_PROMPT


class DesignAnalyzerAgent(BaseAgent):
    name = "DesignAnalyzer"
    description = "Чтение дизайн-проектов (PDF, DWG, DXF, DOCX, Excel)"
    system_prompt = DESIGN_ANALYZER_PROMPT

    async def _process(self, task: Dict[str, Any]) -> Dict[str, Any]:
        file_path = task.get("file_path", "")
        file_content = task.get("file_content", "")

        if file_path and not file_content:
            file_content = await self._extract_content(file_path)

        if not file_content:
            return {"design_analysis": {"error": "Нет содержимого для анализа"}}

        prompt = f"""Проанализируй содержимое проектного документа и извлеки:
1. Список помещений с размерами
2. Виды отделки
3. Спецификацию материалов
4. Инженерные системы

Содержимое документа:
{file_content[:10000]}"""  # Ограничиваем размер

        analysis = await self.ask_llm_json(prompt)
        return {"design_analysis": analysis}

    async def _extract_content(self, file_path: str) -> str:
        """Извлечь текст из файла"""
        ext = os.path.splitext(file_path)[1].lower()

        try:
            if ext == ".pdf":
                return await self._extract_pdf(file_path)
            elif ext in (".dwg", ".dxf"):
                return await self._extract_dxf(file_path)
            elif ext == ".docx":
                return await self._extract_docx(file_path)
            elif ext in (".xlsx", ".xls"):
                return await self._extract_excel(file_path)
            else:
                return f"Неподдерживаемый формат: {ext}"
        except Exception as e:
            self.logger.error(f"Ошибка извлечения из {file_path}: {e}")
            return f"Ошибка чтения файла: {e}"

    async def _extract_pdf(self, path: str) -> str:
        """Извлечь текст из PDF"""
        try:
            import pdfplumber
            text_parts = []
            with pdfplumber.open(path) as pdf:
                for page in pdf.pages[:50]:  # Ограничение страниц
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
                    # Извлекаем таблицы
                    tables = page.extract_tables()
                    for table in tables:
                        for row in table:
                            text_parts.append(" | ".join(str(c) for c in row if c))
            return "\n".join(text_parts)
        except ImportError:
            # Fallback на PyMuPDF
            import fitz
            doc = fitz.open(path)
            text_parts = []
            for page in doc[:50]:
                text_parts.append(page.get_text())
            return "\n".join(text_parts)

    async def _extract_dxf(self, path: str) -> str:
        """Извлечь данные из DXF"""
        import ezdxf
        doc = ezdxf.readfile(path)
        msp = doc.modelspace()

        elements = []
        for entity in msp:
            if entity.dxftype() == "TEXT":
                elements.append(f"Текст: {entity.dxf.text}")
            elif entity.dxftype() == "MTEXT":
                elements.append(f"Текст: {entity.text}")
            elif entity.dxftype() == "LINE":
                start = entity.dxf.start
                end = entity.dxf.end
                length = ((end.x - start.x)**2 + (end.y - start.y)**2)**0.5
                elements.append(f"Линия: длина {length:.2f}")
            elif entity.dxftype() == "LWPOLYLINE":
                points = list(entity.get_points())
                elements.append(f"Полилиния: {len(points)} точек")

        return "\n".join(elements[:500])  # Ограничиваем

    async def _extract_docx(self, path: str) -> str:
        """Извлечь текст из DOCX"""
        from docx import Document
        doc = Document(path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # Извлекаем таблицы
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text for cell in row.cells]
                paragraphs.append(" | ".join(cells))

        return "\n".join(paragraphs)

    async def _extract_excel(self, path: str) -> str:
        """Извлечь данные из Excel"""
        import openpyxl
        wb = openpyxl.load_workbook(path, read_only=True)
        text_parts = []

        for sheet in wb.sheetnames[:5]:
            ws = wb[sheet]
            text_parts.append(f"=== Лист: {sheet} ===")
            for row in ws.iter_rows(max_row=200, values_only=True):
                cells = [str(c) for c in row if c is not None]
                if cells:
                    text_parts.append(" | ".join(cells))

        return "\n".join(text_parts)
